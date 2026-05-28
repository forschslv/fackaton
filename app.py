import os
from dotenv import load_dotenv
# Load environment variables first
load_dotenv()

import json
import csv
import io
import random
import hashlib
import uuid
from datetime import datetime, timezone
from functools import wraps
from flask import Flask, render_template, request, redirect, url_for, flash, session, send_file, abort, jsonify
from werkzeug.utils import secure_filename
from docx import Document

from models import db, User, Product, Order, OrderItem, SystemLog, ProductReview, ChatMessage, InterestGroup

app = Flask(__name__)
app.config['SECRET_KEY'] = os.environ.get('SECRET_KEY', 'marketplace_master_secret_2026_dev_prod')

# Use absolute path for SQLite file to ensure consistency
basedir = os.path.abspath(os.path.dirname(__file__))
app.config['SQLALCHEMY_DATABASE_URI'] = 'sqlite:///' + os.path.join(basedir, 'marketplace.db')
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False

# File upload configuration for payment proofs and bulk imports
UPLOAD_FOLDER = os.path.join(basedir, 'static', 'uploads')
os.makedirs(os.path.join(UPLOAD_FOLDER, 'payment_proofs'), exist_ok=True)
os.makedirs(os.path.join(UPLOAD_FOLDER, 'bulk_imports'), exist_ok=True)
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024  # 16 MB limit

# Initialize db using Flask app context
db.init_app(app)

APP_SETTINGS = {
    "site_name": "MarketHub",
    "currency_symbol": "$",
    "buyer_label": "Покупатель",
    "seller_label": "Продавец",
    "category_1": "Hardware Kits",
    "category_2": "Developer Gear",
    "category_3": "Sponsor API Bundles"
}

@app.context_processor
def inject_settings():
    return dict(cfg=APP_SETTINGS)

# --- MIDDLEWARES & ROUTE DECORATORS ---

def log_system_event(event_type, message):
    """Saves system audit events to database for Admin logs."""
    try:
        new_log = SystemLog(event_type=event_type, message=message, created_at=datetime.now(timezone.utc).replace(tzinfo=None))
        db.session.add(new_log)
        db.session.commit()
    except Exception as e:
        db.session.rollback()
        print(f"Failed to log event: {e}")

def role_required(roles):
    """Decorator to enforce strict user permission access on endpoints and verify identity existence."""
    def decorator(f):
        @wraps(f)
        def decorated_function(*args, **kwargs):
            if 'user_id' not in session:
                flash('Please log in using your credentials to continue.', 'error')
                return redirect(url_for('login'))
            user = User.query.get(session['user_id'])
            if not user:
                session.clear()
                flash('Your session is invalid as the user account was reset or not found. Please log in again.', 'error')
                return redirect(url_for('login'))
            if user.role not in roles or session.get('role') not in roles:
                flash(f'Unauthorized access. This area is restricted to {", ".join(roles)}.', 'error')
                return redirect(url_for('index'))
            return f(*args, **kwargs)
        return decorated_function
    return decorator


@app.after_request
def add_cache_prevent_headers(response):
    """Enforces no-cache headers on all API operations, preventing account switching from leaking cached content in the browser."""
    if request.path.startswith('/api/'):
        response.headers['Cache-Control'] = 'no-store, no-cache, must-revalidate, post-check=0, pre-check=0, max-age=0'
        response.headers['Pragma'] = 'no-cache'
        response.headers['Expires'] = '0'
    return response


@app.context_processor
def inject_global_vars():
    """Injects user authentication status, cart items count and balance globally across templates."""
    current_user = None
    cart_count = 0
    if 'user_id' in session:
        current_user = User.query.get(session['user_id'])
        if not current_user:
            # Session exists but user doesn't (such as after DB reset). Clear gracefully to avoid errors.
            session.clear()
        else:
            # Calculate active cart items count
            cart_order = Order.query.filter_by(buyer_id=session['user_id'], status='Cart').first()
            if cart_order:
                cart_count = sum(item.quantity for item in cart_order.items)
    return dict(current_user=current_user, cart_count=cart_count)


# --- SEEDING COMPILATION UTILITY ---

def run_db_seed():
    """Clears existing database tables, runs schemas instantiation and inserts 15 highly detailed products."""
    print("Beginning database wipe & instantiating clean schemas...")
    db.drop_all()
    db.create_all()
    
    print("Seeding default core users...")
    # 1. Admin
    admin_user = User(full_name="System Administrator", email="admin@marketplace.com", role="Admin", balance=1000.0)
    admin_user.set_password("12345678A")
    # 2. Seller
    seller_user = User(full_name="TechVanguard Merchant", email="seller@marketplace.com", role="Seller", balance=500.0)
    seller_user.set_password("12345678A")
    # 3. Buyer
    buyer_user = User(full_name="Alex Developer", email="buyer@marketplace.com", role="Buyer", balance=300.0)
    buyer_user.set_password("12345678A")
    
    db.session.add_all([admin_user, seller_user, buyer_user])
    db.session.commit()
    
    # Reload seller to link products
    seller = User.query.filter_by(role="Seller").first()
    
    # 15 Highly Detailed Products partitioned across 3 Categories
    sample_listings = [
        # --- CATEGORY A: Hardware Kits ---
        {
            "title": "IoT Starter Pack Pro",
            "description": "Unlock high-performance IoT development with this comprehensive sensor expansion kit. Perfect for hackers looking to aggregate environmental telemetry or build automated smart-office prototypes.",
            "category": "Hardware Kits",
            "price": 59.99,
            "stock": 25,
            "image_url": "https://images.unsplash.com/photo-1518770660439-4636190af475?auto=format&fit=crop&w=600&q=80",
            "metadata": {"difficulty": "Beginner", "mcu": "ESP32-S3", "sensors": "Temp, Humidity, Light", "warranty": "6 Months"}
        },
        {
            "title": "Sensory Node Array",
            "description": "Industrial grade mesh sensing cluster. Designed to dispatch localized environmental specifications over securely encrypted Zigbee 3.0 protocol configurations.",
            "category": "Hardware Kits",
            "price": 129.99,
            "stock": 10,
            "image_url": "https://images.unsplash.com/photo-1555664424-778a1e5e1b48?auto=format&fit=crop&w=600&q=80",
            "metadata": {"protocol": "Zigbee 3.0", "range": "100m", "power": "LiPo 1200mAh Included", "accuracy": "+/- 0.05%"}
        },
        {
            "title": "Battery Shield V2",
            "description": "Heavy-duty power management board configured to protect prototype boards from surge anomalies while ensuring continuous stable delivery over custom intervals.",
            "category": "Hardware Kits",
            "price": 19.99,
            "stock": 50,
            "image_url": "https://images.unsplash.com/photo-1563770660941-20978e870e26?auto=format&fit=crop&w=600&q=80",
            "metadata": {"chemistry": "Li-Ion", "output": "5V @ 2A", "charging": "USB-C Fast Charging", "protection": "Overload/Short Circuit"}
        },
        {
            "title": "OLED Dual Display Shield",
            "description": "Vivid dual-layer high refresh display shield supporting pixel alignment algorithms for real-time diagnostic graphics rendering.",
            "category": "Hardware Kits",
            "price": 34.50,
            "stock": 15,
            "image_url": "https://images.unsplash.com/photo-1558494949-ef010cbdcc31?auto=format&fit=crop&w=600&q=80",
            "metadata": {"resolution": "128x64 x2 Nodes", "interface": "I2C / SPI", "size": "0.96 inch", "driver": "SSD1306"}
        },
        {
            "title": "Robotic Arm Actuators Pack",
            "description": "Surgical grade precise feedback actuator kit consisting of high torque metal gear servos, mounting plates, and dynamic torque calibration modules.",
            "category": "Hardware Kits",
            "price": 249.00,
            "stock": 5,
            "image_url": "https://images.unsplash.com/photo-1485827404703-89b55fcc595e?auto=format&fit=crop&w=600&q=80",
            "metadata": {"servos": "4x SG90 + 2x MG996R", "payload_max": "150g", "voltage": "6V to 8.4V", "chassis": "Anodized Aluminum"}
        },
        # --- CATEGORY B: Developer Gear ---
        {
            "title": "Retro Mech Keycaps Set",
            "description": "Double-shot PBT keycap layout featuring an old-school aesthetic. Sculpted to withstand exhaustive developer terminal operations with clear legibility and zero fading.",
            "category": "Developer Gear",
            "price": 45.00,
            "stock": 30,
            "image_url": "https://images.unsplash.com/photo-1587829741301-dc798b83add3?auto=format&fit=crop&w=600&q=80",
            "metadata": {"profile": "Cherry Profile", "material": "Double-shot PBT", "keys_count": "135 keys", "compatibility": "ANSI Layout MX"}
        },
        {
            "title": "Ergonomic Split Mechanical Keyboard",
            "description": "Revolutionary split ergonomic ortholinear layout reducing wrist strain. Standardized with hot-swappable switches and dual-layer layout customization.",
            "category": "Developer Gear",
            "price": 189.99,
            "stock": 8,
            "image_url": "https://images.unsplash.com/photo-1595225476474-87563907a212?auto=format&fit=crop&w=600&q=80",
            "metadata": {"switches": "Gateron Brown Linear", "layout": "60% Ortholinear Split", "connection": "USB-C / Bluetooth 5.1", "backlight": "RGB"}
        },
        {
            "title": "Pro ANC Developer Headphones",
            "description": "Deep hybrid industrial ANC filters configured to neutralize noisy coffee shop or open space offices, giving developer minds immaculate quietness.",
            "category": "Developer Gear",
            "price": 299.99,
            "stock": 12,
            "image_url": "https://images.unsplash.com/photo-1505740420928-5e560c06d30e?auto=format&fit=crop&w=600&q=80",
            "metadata": {"anc_depth": "40dB Hybrid", "battery_life": "40 Hours", "audio_codec": "LDAC / AAC", "weight": "250g"}
        },
        {
            "title": "Vertical Wireless Ergonomic Mouse",
            "description": "Perfect right-handed neutral handshake angle mouse mapping physical strain to zero levels. Instant variable tracking DPI controller.",
            "category": "Developer Gear",
            "price": 79.50,
            "stock": 20,
            "image_url": "https://images.unsplash.com/photo-1615663245857-ac93bb7c39e7?auto=format&fit=crop&w=600&q=80",
            "metadata": {"sensor": "Optical 4000 DPI", "hand": "Right-handed", "weight": "115g", "charging": "USB-C"}
        },
        {
            "title": "Vegan Leather Heated Desk Mat",
            "description": "Large elegant charcoal mat keeping developer workspace hands toasty during cold programming sprint nights. Equipped with built-in auto-shutdown features.",
            "category": "Developer Gear",
            "price": 39.00,
            "stock": 15,
            "image_url": "https://images.unsplash.com/photo-1632292224971-0d45778bd364?auto=format&fit=crop&w=600&q=80",
            "metadata": {"temp_levels": "3 Settings (35/45/55C)", "timer": "4hr Automatic Shutdown", "dimensions": "80cm x 40cm", "material": "Vegan Leather"}
        },
        # --- CATEGORY C: Sponsor API Bundles ---
        {
            "title": "Multi-Model Premium LLM API Voucher",
            "description": "Activate bulk token allocations across standard state-of-the-art vision, text, and multimodal reasoning API integrations without latency constraints.",
            "category": "Sponsor API Bundles",
            "price": 150.00,
            "stock": 100,
            "image_url": "https://images.unsplash.com/photo-1677442136019-21780efad99a?auto=format&fit=crop&w=600&q=80",
            "metadata": {"tokens": "20 Million Tokens", "validity": "6 Months", "endpoints": "Universal REST/gRPC SDKs", "security": "IP Bound"}
        },
        {
            "title": "Enterprise Mapping Token",
            "description": "Developer entitlement pass providing access to high accuracy street grid APIs, coordinate validation routing tables, and vector tile loaders.",
            "category": "Sponsor API Bundles",
            "price": 85.00,
            "stock": 50,
            "image_url": "https://images.unsplash.com/photo-1524661135-423995f22d0b?auto=format&fit=crop&w=600&q=80",
            "metadata": {"queries": "50,000 requests", "features": "3D Terrain + Address Validation", "supports": "React / IOS / Android SDK", "sla": "99.9% Uptime"}
        },
        {
            "title": "Real-time Vector DB License Bundle",
            "description": "Fully managed distributed cloud storage configured specifically to host embeddings arrays. Instant cosine matching with sub-millisecond retrieval.",
            "category": "Sponsor API Bundles",
            "price": 210.00,
            "stock": 40,
            "image_url": "https://images.unsplash.com/photo-1544383835-bda2bc66a55d?auto=format&fit=crop&w=600&q=80",
            "metadata": {"vectors_limit": "5 Million Embeddings", "max_dimensions": "1536 dims", "tier": "Developer Plus Plan", "sla": "99.95% Server SLA"}
        },
        {
            "title": "Sponsor Cloud Credits $500",
            "description": "Fast-track your database clustering and microservice containers using official platforms. Highly recommended for hackathon project final evaluations.",
            "category": "Sponsor API Bundles",
            "price": 400.00,
            "stock": 30,
            "image_url": "https://images.unsplash.com/photo-1454165804606-c3d57bc86b40?auto=format&fit=crop&w=600&q=80",
            "metadata": {"value": "$500 USD credits", "region": "Global multi-cluster", "restrictions": "Valid for VM, DB, Serverless", "code_expiry": "12 Months"}
        },
        {
            "title": "Security Scanner Premium Dev Tier",
            "description": "Automated security scanning pass checks your code lines against known configuration vulnerabilities, leak points, and compliance policies.",
            "category": "Sponsor API Bundles",
            "price": 120.00,
            "stock": 25,
            "image_url": "https://images.unsplash.com/photo-1526374965328-7f61d4dc18c5?auto=format&fit=crop&w=600&q=80",
            "metadata": {"repos": "Up to 5 Projects", "scans": "Unlimited Real-time Scans", "features": "SAST + DAST Analyzers", "hooks": "GitHub Actions / GitLab Webhooks"}
        }
    ]

    for item in sample_listings:
        prod = Product(
            seller_id=seller.id,
            title=item["title"],
            description=item["description"],
            category=item["category"],
            price=item["price"],
            stock=item["stock"],
            image_url=item["image_url"],
            is_active=True
        )
        prod.set_metadata(item["metadata"])
        db.session.add(prod)
    
    db.session.commit()
    
    # Seed default interest groups
    default_groups = [
        {
            "slug": "group_dev",
            "name": "💻 Разработка и Код",
            "description": "Обсуждение программирования, железа и IT-технологий"
        },
        {
            "slug": "group_marketing",
            "name": "📈 Маркетинг и SEO",
            "description": "Продвижение продуктов, реклама и SEO-оптимизация"
        },
        {
            "slug": "group_design",
            "name": "🎨 Дизайн и UI/UX",
            "description": "Креатив, проектирование интерфейсов и UX-аналитика"
        }
    ]
    for g in default_groups:
        group = InterestGroup(slug=g["slug"], name=g["name"], description=g["description"])
        db.session.add(group)
    db.session.commit()
    
    log_system_event("SYSTEM", "Database successfully cleaned and loaded with 15 premium sample listings.")
    print("Marketplace database initialization complete! Seeded Admin, Seller, Buyer, and 15 product assets.")


@app.cli.command("init-db")
def init_db_command():
    """Wipes database schemas and generates initial boilerplate datasets."""
    run_db_seed()
    print("Bootstrap configurations generated successfully.")


# Database auto-builder/initializer for runtime safety
# @app.before_all_handler if hasattr(app, 'before_all_handler') else None
# def auto_init_database():
#     pass

# Ensure tables are built and seeded automatically for immediate viewing in browser
with app.app_context():
    db.create_all()
    
    # Dynamic seed fallback for newly introduced interest groups table
    try:
        if InterestGroup.query.count() == 0:
            default_groups = [
                {"slug": "group_dev", "name": "💻 Разработка и Код", "description": "Обсуждение программирования, железа и IT-технологий"},
                {"slug": "group_marketing", "name": "📈 Маркетинг и SEO", "description": "Продвижение продуктов, реклама и SEO-оптимизация"},
                {"slug": "group_design", "name": "🎨 Дизайн и UI/UX", "description": "Креатив, проектирование интерфейсов и UX-аналитика"}
            ]
            for g in default_groups:
                group = InterestGroup(slug=g["slug"], name=g["name"], description=g["description"])
                db.session.add(group)
            db.session.commit()
            print("Successfully bootstrapped default interest chats inside existing database.")
    except Exception as ex:
        print(f"Error checking/seeding interest groups on startup: {ex}")

    if not os.path.exists(os.path.join(basedir, 'marketplace.db')):
        run_db_seed()


# --- PUBLIC VIEWS & GENERAL MODULES ---

@app.route('/')
def index():
    """Executive Landing Page with Dynamic Showcase."""
    # Group categories
    categories = db.session.query(Product.category).distinct().all()
    categories = [c[0] for c in categories]
    
    featured_products = Product.query.filter_by(is_active=True).limit(4).all()
    
    return render_template('index.html', categories=categories, featured_products=featured_products)


@app.route('/catalog')
def catalog():
    """Advanced Filtering, Text querying and Dynamic Sorting Catalog Page."""
    query = request.args.get('q', '').strip()
    category = request.args.get('category', '').strip()
    sort_by = request.args.get('sort', '').strip()
    
    # Base Query
    p_query = Product.query.filter_by(is_active=True)
    
    # Category Filter
    if category:
        p_query = p_query.filter_by(category=category)
    
    # Text Search Filter
    if query:
        p_query = p_query.filter(
            (Product.title.like(f"%{query}%")) | 
            (Product.description.like(f"%{query}%"))
        )
    
    # Sorting
    if sort_by == 'price_asc':
        p_query = p_query.order_by(Product.price.asc())
    elif sort_by == 'price_desc':
        p_query = p_query.order_by(Product.price.desc())
    elif sort_by == 'stock_desc':
        p_query = p_query.order_by(Product.stock.desc())
    else:
        p_query = p_query.order_by(Product.id.desc())
        
    products = p_query.all()
    
    all_categories = db.session.query(Product.category).distinct().all()
    all_categories = [c[0] for c in all_categories]
    
    return render_template('catalog.html', products=products, categories=all_categories, active_category=category, search_query=query, sort_by=sort_by)


@app.route('/api/ai-consultant/<int:product_id>')
def ai_consultant(product_id):
    """Generates detailed, personalized product review & hackathon integration consultation using Gemini AI."""
    product = Product.query.get_or_404(product_id)
    
    api_key = os.environ.get("GEMINI_API_KEY")
    if not api_key:
        return jsonify({
            "success": False,
            "error": "Ключ API (GEMINI_API_KEY) не настроен в настройках (Secrets). Пожалуйста, укажите его для включения ИИ-советника."
        }), 400
        
    try:
        import google.generativeai as genai
        
        genai.configure(api_key=api_key)
        # Use a modern, non-deprecated fast model
        model = genai.GenerativeModel('gemini-2.5-flash')
        
        prompt = f"""
        Ты — ведущий технологический советник, эксперт хакатонов и опытный бизнес-консультант.
        Твоя задача — проанализировать товар "{product.title}" и дать краткие, но невероятно убедительные рекомендации по его интеграции и использованию.
        
        Категория товара: {product.category}
        Цена: {APP_SETTINGS['currency_symbol']}{product.price}
        Текущий запас на складе: {product.stock} шт.
        Описание товара: {product.description}
        
        Пожалуйста, сформируй структурированный разбор на русском языке, состоящий из следующих разделов (сделай оформление в красивом Markdown):
        1. **🚀 Ценность для Хакатона** — как этот товар может ускорить разработку, стать киллер-фичей MVP или улучшить защиту проекта перед жюри.
        2. **💡 Сценарии интеграции** — конкретные технические кодовые примеры или архитектурные сценарии, как разработчики могут внедрить его.
        3. **🎯 Целевая аудитория** — кому этот продукт необходим в первую очередь.
        
        Пиши профессионально, вдохновляюще, лаконично и структурно.
        """
        
        response = model.generate_content(prompt)
        return jsonify({
            "success": True,
            "analysis": response.text
        })
    except Exception as e:
        print(f"Gemini API Exception: {e}")
        return jsonify({
            "success": False,
            "error": f"Произошла ошибка при обращении к ИИ-ассистенту: {str(e)}"
        }), 500


@app.route('/product/<int:product_id>')
def product_detail(product_id):
    """Dynamic High-Fidelity Product Detail screen reflecting customized specifications from metadata_json."""
    product = Product.query.get_or_404(product_id)
    return render_template('product_detail.html', product=product)


# --- SECURITY & REGISTRATION MODULES ---

def dispatch_2fa_otp(user, otp_code):
    """Dispatches the OTP code via the method specified in config.json."""
    config_path = os.path.join(basedir, 'config.json')
    delivery_method = "console"
    smtp_server = "smtp.gmail.com"
    smtp_port = 587
    smtp_username = ""
    smtp_password = ""
    smtp_sender = "no-reply@marketplace.com"
    
    if os.path.exists(config_path):
        try:
            with open(config_path, 'r') as f:
                cfg = json.load(f)
                delivery_method = cfg.get("2fa_delivery_method", "console").strip().lower()
                smtp_server = cfg.get("smtp_server", smtp_server)
                smtp_port = int(cfg.get("smtp_port", smtp_port))
                smtp_username = cfg.get("smtp_username", smtp_username)
                smtp_password = cfg.get("smtp_password", smtp_password)
                smtp_sender = cfg.get("smtp_sender", smtp_sender)
        except Exception as e:
            print(f"Error reading config.json: {e}")
            
    # Fallback/Override with env variables if present
    smtp_username = os.environ.get("SMTP_USERNAME") or smtp_username
    smtp_password = os.environ.get("SMTP_PASSWORD") or smtp_password
    smtp_sender = os.environ.get("SMTP_SENDER") or smtp_sender
    
    if delivery_method == "email" and smtp_username and smtp_password:
        import smtplib
        from email.mime.text import MIMEText
        from email.mime.multipart import MIMEMultipart
        
        try:
            msg = MIMEMultipart()
            msg['From'] = smtp_sender
            msg['To'] = user.email
            msg['Subject'] = f"{otp_code} is your Security OTP Key"
            
            body = (
                f"Hello {user.full_name},\n\n"
                f"Your 2FA verification challenge code is: {otp_code}\n\n"
                f"Input this code on the verification portal to log in.\n\n"
                f"If you did not make this request, please ignore this email.\n\n"
                f"Best regards,\n"
                f"MarketHub Team"
            )
            msg.attach(MIMEText(body, 'plain'))
            
            # Send the email
            if smtp_port == 465:
                server = smtplib.SMTP_SSL(smtp_server, smtp_port, timeout=10)
            else:
                server = smtplib.SMTP(smtp_server, smtp_port, timeout=10)
                server.starttls()
                
            server.login(smtp_username, smtp_password)
            server.sendmail(smtp_sender, user.email, msg.as_string())
            server.quit()
            
            log_system_event("AUTH", f"Successfully dispatched 2FA mail to {user.email}")
            return "email"
        except Exception as e:
            err_msg = f"Failed sending real 2FA email to {user.email}: {e}"
            print(err_msg)
            log_system_event("AUTH", err_msg)
            
            # Fallback to console printing on SMTP failure
            print("\n" + "="*70)
            print(f" [OTP SECURITY CODE CHALLENGE (FALLBACK due to SMTP Error)] ")
            print(f" Recipient: {user.full_name} ({user.email})")
            print(f" Security OTP Challenge Response Code: {otp_code}")
            print("="*70 + "\n", flush=True)
            return "console_fallback"
    else:
        # Standard Terminal print
        print("\n" + "="*70)
        print(f" [OTP SECURITY CODE CHALLENGE] ")
        print(f" Recipient: {user.full_name} ({user.email})")
        print(f" Security OTP Challenge Response Code: {otp_code}")
        print("="*70 + "\n", flush=True)
        return "console"


@app.route('/register', methods=['GET', 'POST'])
def register():
    """Secure Customer & Merchant Enrollment Portal."""
    if request.method == 'POST':
        full_name = request.form.get('full_name').strip()
        email = request.form.get('email').strip().lower()
        password = request.form.get('password')
        role = request.form.get('role', 'Buyer').strip()
        two_factor_opt = request.form.get('two_factor_enabled') # 'on' or None
        
        two_factor_enabled = (two_factor_opt == 'on')
        
        if role not in ['Buyer', 'Seller']:
            role = 'Buyer'
            
        if not full_name or not email or not password:
            flash("All registration fields are required.", "error")
            return redirect(url_for('register'))
            
        existing_user = User.query.filter_by(email=email).first()
        if existing_user:
            flash("Email address is already in use onto our network.", "error")
            return redirect(url_for('register'))
            
        new_user = User(
            full_name=full_name,
            email=email,
            role=role,
            balance=100.0 if role == 'Buyer' else 0.0,  # Give $100 starter tokens to buyers for convenience
            two_factor_enabled=two_factor_enabled
        )
        new_user.set_password(password)
        db.session.add(new_user)
        db.session.commit()
        
        log_system_event("USER", f"New user '{full_name}' successfully registered as a {role} matching email {email}. 2FA enabled: {two_factor_enabled}")
        flash("Registration successful! You may now sign in using your credentials.", "success")
        return redirect(url_for('login'))
        
    return render_template('register.html')


@app.route('/login', methods=['GET', 'POST'])
def login():
    """Standard Credentials Credentials Auth Route initiating dynamic OTP delivery."""
    if request.method == 'POST':
        email = request.form.get('email').strip().lower()
        password = request.form.get('password')
        
        user = User.query.filter_by(email=email).first()
        if not user or not user.check_password(password):
            flash("Incorrect credentials. Please verify your info and retry.", "error")
            return redirect(url_for('login'))
            
        # Check if 2FA is active/enabled for this user
        if not user.two_factor_enabled:
            # Transfer authorization state and log in immediately
            session.clear()
            session['user_id'] = user.id
            session['user_name'] = user.full_name
            session['role'] = user.role
            
            log_system_event("AUTH", f"User logged in successfully (2FA disabled): {user.email}")
            flash(f"Welcome back, {user.full_name}! (Signed in successfully, 2FA bypassed)", "success")
            
            if user.role == 'Admin':
                return redirect(url_for('admin_dashboard'))
            elif user.role == 'Seller':
                return redirect(url_for('seller_dashboard'))
            else:
                return redirect(url_for('catalog'))
                
        # Initiate 2FA OTP Delivery
        otp_code = "".join([str(random.randint(0, 9)) for _ in range(6)])
        
        # Buffer code in session state
        session['temp_otp'] = otp_code
        session['temp_otp_user_id'] = user.id
        session['temp_otp_expires_at'] = datetime.now(timezone.utc).replace(tzinfo=None).isoformat()
        
        # Dispatch code dynamically using config.json specifications
        dispatch_method = dispatch_2fa_otp(user, otp_code)
        
        log_system_event("AUTH", f"Generated 2FA challenge of {otp_code} for user: {user.email} (Dispatched via {dispatch_method})")
        if dispatch_method == "email":
            flash("Standard password accepted. A secure 2FA challenge has been emailed to you.", "info")
        else:
            flash("Standard password accepted. A secure 6-digit OTP code has been dispatched to your terminal.", "info")
            
        return redirect(url_for('verify_otp'))
        
    return render_template('login.html')


@app.route('/verify-otp', methods=['GET', 'POST'])
def verify_otp():
    """2FA Challenge Verification Terminal input page."""
    if 'temp_otp_user_id' not in session or 'temp_otp' not in session:
        flash("Session invalid. Realignment required.", "error")
        return redirect(url_for('login'))
        
    if request.method == 'POST':
        entered_otp = request.form.get('otp', '').strip()
        expected_otp = session.get('temp_otp')
        
        if entered_otp == expected_otp:
            # Code matches, fully log in
            user_id = session.get('temp_otp_user_id')
            user = User.query.get(user_id)
            
            # Transfer authorization state
            session.clear()
            session['user_id'] = user.id
            session['user_name'] = user.full_name
            session['role'] = user.role
            
            log_system_event("AUTH", f"User logged in successfully via OTP verification verification: {user.email}")
            flash(f"Welcome back, {user.full_name}! Security authentication approved.", "success")
            
            # Destination redirects targeting workspace divisions
            if user.role == 'Admin':
                return redirect(url_for('admin_dashboard'))
            elif user.role == 'Seller':
                return redirect(url_for('seller_dashboard'))
            else:
                return redirect(url_for('index'))
        else:
            flash("Invalid security OTP verification challenge code. Please retry.", "error")
            return redirect(url_for('verify_otp'))
            
    return render_template('verify_otp.html')


@app.route('/logout')
def logout():
    """Invalides session and returns to Landing page."""
    if 'user_name' in session:
        log_system_event("AUTH", f"Session destroyed on demand of user_id: {session.get('user_id')}")
    session.clear()
    flash("Session terminated successfully.", "success")
    return redirect(url_for('index'))


@app.route('/settings', methods=['GET', 'POST'])
def settings():
    """Manage User Profiles & Account Security Credentials."""
    if 'user_id' not in session:
        flash("Please sign in to access settings.", "error")
        return redirect(url_for('login'))
        
    user = User.query.get(session['user_id'])
    if not user:
        session.clear()
        return redirect(url_for('login'))
        
    if request.method == 'POST':
        full_name = request.form.get('full_name', '').strip()
        new_password = request.form.get('new_password', '').strip()
        two_factor_opt = request.form.get('two_factor_enabled') # 'on' or None
        
        two_factor_enabled = (two_factor_opt == 'on')
        
        if not full_name:
            flash("Profile name cannot be empty.", "error")
            return redirect(url_for('settings'))
            
        user.full_name = full_name
        user.two_factor_enabled = two_factor_enabled
        
        if new_password:
            user.set_password(new_password)
            
        try:
            db.session.commit()
            log_system_event("USER", f"User {user.email} updated profile parameters/settings. 2FA is now: {two_factor_enabled}")
            flash("Account settings successfully saved!", "success")
            return redirect(url_for('settings'))
        except Exception as e:
            db.session.rollback()
            flash(f"Error during settings save configuration: {e}", "error")
            return redirect(url_for('settings'))
            
    return render_template('settings.html', user=user)


# --- MESSENGER & PRODUCT REVIEWS MODULES ---

@app.route('/messenger')
def messenger():
    """Universal Messenger Workstation featuring Chat rooms, PM and Support divisions."""
    if 'user_id' not in session:
        flash("Пожалуйста, войдем в аккаунт, чтобы открыть мессенджер.", "error")
        return redirect(url_for('login'))
        
    current_uid = session['user_id']
    user = User.query.get(current_uid)
    if not user:
        session.clear()
        return redirect(url_for('login'))
        
    # Get all other users for direct chat initiation
    all_users = User.query.filter(User.id != current_uid).all()
    
    # Check if a direct chat room was requested from product details
    target_partner_id = request.args.get('start_chat_with', type=int)
    preselected_room = "lobby"
    if target_partner_id:
        partner = User.query.get(target_partner_id)
        if partner:
            # Deterministic room_id for DM between current_uid and partner_id
            u1, u2 = sorted([current_uid, target_partner_id])
            preselected_room = f"dm_{u1}_{u2}"
            
    # Also support general room preselection from Query parameters
    room_param = request.args.get('room', '').strip()
    if room_param:
        preselected_room = room_param
        
    return render_template('messenger.html', 
                           all_users=all_users, 
                           preselected_room=preselected_room)


@app.route('/api/messenger/contacts')
def api_messenger_contacts():
    """Retrieves contacts list categorized by Support, Direct Messages, and Public Channels."""
    if 'user_id' not in session:
        return jsonify({'error': 'Unauthorized'}), 401
    
    current_uid = session['user_id']
    current_user = User.query.get(current_uid)
    if not current_user:
        return jsonify({'error': 'Unauthorized'}), 401
        
    contacts = []
    
    # 1. Broadly Available General Lobby Chat
    contacts.append({
        'id': 'lobby',
        'name': '💬 Общий чат покупателей (Lobby)',
        'type': 'group',
        'role_badge': 'Публичный',
        'subtitle': 'Открытый канал для всех участников',
        'avatar_char': 'L',
        'partner_id': None
    })
    
    # Groups by interest loaded dynamically from database
    try:
        db_groups = InterestGroup.query.order_by(InterestGroup.id.asc()).all()
        for g in db_groups:
            # Extract first character or emoji
            char = g.name[0] if g.name else '#'
            contacts.append({
                'id': f"group_{g.slug}",
                'name': g.name,
                'type': 'group',
                'role_badge': 'Интересы',
                'subtitle': g.description or 'Группа по интересам',
                'avatar_char': char,
                'partner_id': None,
                'slug': g.slug,
                'created_by_id': g.created_by_id
            })
    except Exception as ex:
        print(f"Error loading interest groups: {ex}")
        contacts.append({
            'id': 'group_dev',
            'name': '💻 Разработка и Код',
            'type': 'group',
            'role_badge': 'Интересы',
            'subtitle': 'Обсуждение программирования',
            'avatar_char': '💻',
            'partner_id': None
        })
    
    # 2. Support channels view
    if current_user.role == 'Admin':
        # Admin gets a list of all custom support channels that have messages
        support_rooms = db.session.query(ChatMessage.room_id).filter(
            ChatMessage.room_id.like('support_%')
        ).distinct().all()
        
        for (room_id,) in support_rooms:
            try:
                buyer_id = int(room_id.split('_')[1])
                buyer_user = User.query.get(buyer_id)
                if buyer_user:
                    contacts.append({
                        'id': room_id,
                        'name': f"🛠️ Поддержка: {buyer_user.full_name}",
                        'type': 'support',
                        'role_badge': buyer_user.role,
                        'subtitle': f"Запрос от {buyer_user.email}",
                        'avatar_char': buyer_user.full_name[0].upper(),
                        'partner_id': buyer_user.id
                    })
            except Exception:
                continue
    else:
        # Standard user gets a single persistent support channel to Admin
        admin_user = User.query.filter_by(role='Admin').first()
        contacts.append({
            'id': f"support_{current_uid}",
            'name': '🛠️ Круглосуточная служба поддержки',
            'type': 'support',
            'role_badge': 'Поддержка',
            'subtitle': 'Связь с администрацией площадки',
            'avatar_char': 'S',
            'partner_id': admin_user.id if admin_user else None
        })

    # 3. Direct Message Partners from past message logs
    # Past message rooms that start with "dm_" and contain the user's ID
    dm_rooms_query = db.session.query(ChatMessage.room_id).filter(
        ChatMessage.room_id.like('dm_%')
    ).distinct().all()
    
    active_partners = set()
    for (room_id,) in dm_rooms_query:
        parts = room_id.split('_')
        if len(parts) == 3:
            try:
                u1, u2 = int(parts[1]), int(parts[2])
                if current_uid == u1:
                    active_partners.add(u2)
                elif current_uid == u2:
                    active_partners.add(u1)
            except ValueError:
                continue
                
    for partner_id in active_partners:
        partner = User.query.get(partner_id)
        if partner:
            u1, u2 = sorted([current_uid, partner_id])
            contacts.append({
                'id': f"dm_{u1}_{u2}",
                'name': partner.full_name,
                'type': 'direct',
                'role_badge': 'Продавец' if partner.role == 'Seller' else 'Покупатель',
                'subtitle': partner.email,
                'avatar_char': partner.full_name[0].upper(),
                'partner_id': partner.id
            })
            
    return jsonify({'contacts': contacts})


@app.route('/api/messenger/messages')
def api_get_messages():
    """Retrieves messages in a specific room, with sequential filtering after_id."""
    if 'user_id' not in session:
        return jsonify({'error': 'Unauthorized'}), 401
        
    room_id = request.args.get('room_id', 'lobby').strip()
    after_id = request.args.get('after_id', type=int)
    
    # security guard: make sure users only read rooms they are authorized to hold
    current_uid = session['user_id']
    current_role = session['role']
    
    if room_id != 'lobby':
        if room_id.startswith('support_'):
            try:
                owner_id = int(room_id.split('_')[1])
                if current_role != 'Admin' and current_uid != owner_id:
                    return jsonify({'error': 'Forbidden access to support line'}), 403
            except ValueError:
                return jsonify({'error': 'Malformed room_id'}), 400
        elif room_id.startswith('dm_'):
            try:
                parts = room_id.split('_')
                u1, u2 = int(parts[1]), int(parts[2])
                if current_uid != u1 and current_uid != u2:
                    return jsonify({'error': 'Forbidden access to direct message room'}), 403
            except ValueError:
                return jsonify({'error': 'Malformed room_id'}), 400
        elif room_id.startswith('group_'):
            pass # Open to all users
        else:
            return jsonify({'error': 'Malformed room_id'}), 400

    query = ChatMessage.query.filter_by(room_id=room_id)
    if after_id:
        query = query.filter(ChatMessage.id > after_id)
        
    messages = query.order_by(ChatMessage.id.asc()).all()
    
    return jsonify({
        'messages': [msg.to_dict() for msg in messages]
    })


@app.route('/api/messenger/send', methods=['POST'])
def api_send_message():
    """Dispatches a message onto a designated chat room with optional image attachments."""
    if 'user_id' not in session:
        return jsonify({'error': 'Unauthorized'}), 401
        
    current_uid = session['user_id']
    room_id = request.form.get('room_id', 'lobby').strip()
    message_text = request.form.get('message', '').strip()
    receiver_id = request.form.get('receiver_id', type=int)
    
    # Handle File/Attachment Upload inside chat with robust unique naming!
    attachment_url = None
    file = request.files.get('attachment')
    if file and file.filename != '':
        _, ext = os.path.splitext(file.filename)
        if not ext:
            ext = '.png'
        unique_name = f"chat_{uuid.uuid4().hex}{ext.lower()}"
        os.makedirs(os.path.join(app.config['UPLOAD_FOLDER'], 'chat_attachments'), exist_ok=True)
        save_path = os.path.join(app.config['UPLOAD_FOLDER'], 'chat_attachments', unique_name)
        file.save(save_path)
        attachment_url = f"/static/uploads/chat_attachments/{unique_name}"
        
    if not message_text and not attachment_url:
        return jsonify({'error': 'Cannot post blank message.'}), 400
        
    # Standardize support channels on the client
    if room_id.startswith('support_'):
        pass # Already computed room
        
    # Create the DB message record
    new_message = ChatMessage(
        sender_id=current_uid,
        receiver_id=receiver_id,
        room_id=room_id,
        message=message_text,
        attachment_url=attachment_url,
        created_at=datetime.now(timezone.utc).replace(tzinfo=None)
    )
    
    db.session.add(new_message)
    db.session.commit()
    
    return jsonify({
        'status': 'ok',
        'message': new_message.to_dict()
    })


@app.route('/api/messenger/delete/<int:message_id>', methods=['POST'])
@role_required(['Admin'])
def api_delete_message(message_id):
    """Admin endpoint to moderate and delete messages."""
    msg = ChatMessage.query.get(message_id)
    if not msg:
        return jsonify({'error': 'Message not found'}), 404
    
    db.session.delete(msg)
    db.session.commit()
    log_system_event("MODERATION", f"Admin {session['user_id']} deleted message {message_id} from room {msg.room_id}")
    return jsonify({'status': 'ok'})


@app.route('/api/messenger/create_group', methods=['POST'])
def api_create_group():
    """Allows users to register interest groups in the messenger workstation."""
    if 'user_id' not in session:
        return jsonify({'error': 'Unauthorized'}), 401
    
    name = request.form.get('name', '').strip()
    description = request.form.get('description', '').strip()
    
    if not name:
        return jsonify({'error': 'Название группы обязательно.'}), 400
        
    import re
    # Generate slug
    base_slug = re.sub(r'[^a-zA-Z0-9]', '_', name.lower())
    if not base_slug or base_slug == '_':
        base_slug = "interest"
        
    slug = f"custom_{base_slug}"
    count = 1
    original_slug = slug
    while InterestGroup.query.filter_by(slug=slug).first():
        slug = f"{original_slug}_{count}"
        count += 1
        
    new_group = InterestGroup(
        slug=slug,
        name=name,
        description=description,
        created_by_id=session['user_id']
    )
    db.session.add(new_group)
    db.session.commit()
    
    log_system_event("CHAT", f"Created new interest group: '{name}' (slug: {slug}) by user {session['user_id']}")
    
    return jsonify({
        'status': 'ok',
        'group': new_group.to_dict()
    })


@app.route('/api/messenger/delete_group/<string:slug>', methods=['POST'])
def api_delete_group(slug):
    """Allows group creators or admins to retire specific interest channels."""
    if 'user_id' not in session:
        return jsonify({'error': 'Unauthorized'}), 401
        
    group = InterestGroup.query.filter_by(slug=slug).first()
    if not group:
        return jsonify({'error': 'Группа не найдена.'}), 404
        
    current_uid = session['user_id']
    current_role = session['role']
    
    if current_role != 'Admin' and group.created_by_id != current_uid:
        return jsonify({'error': 'Недостаточно прав. Только создатель или администратор может удалить эту группу.'}), 403
        
    room_id = f"group_{slug}"
    # Clean up associated chat logs
    ChatMessage.query.filter_by(room_id=room_id).delete()
    
    db.session.delete(group)
    db.session.commit()
    
    log_system_event("CHAT", f"Deleted interest group '{group.name}' and all messages in room {room_id}")
    
    return jsonify({'status': 'ok'})

# --- PRODUCT REVIEWS CONTROLLER ACTIONS ---

@app.route('/product/<int:product_id>/review', methods=['POST'])
def submit_review(product_id):
    """Enables users to leave dynamic rating comments on active catalog listings."""
    if 'user_id' not in session:
        flash("Вам необходимо войти в систему, чтобы оставить отзыв.", "error")
        return redirect(url_for('login'))
        
    rating = request.form.get('rating', type=int)
    comment = request.form.get('comment', '').strip()
    
    if not rating or rating < 1 or rating > 5:
        flash("Ошибка: Укажите оценку от 1 до 5 звезд.", "error")
        return redirect(url_for('product_detail', product_id=product_id))
        
    if not comment:
        flash("Ошибка: Текст отзыва не может быть пустым.", "error")
        return redirect(url_for('product_detail', product_id=product_id))
        
    # Check if user has already reviewed the product
    existing_review = ProductReview.query.filter_by(product_id=product_id, user_id=session['user_id']).first()
    if existing_review:
        existing_review.rating = rating
        existing_review.comment = comment
        existing_review.created_at = datetime.now(timezone.utc).replace(tzinfo=None)
        db.session.commit()
        flash("Ваш предыдущий отзыв был успешно обновлен!", "success")
    else:
        new_review = ProductReview(
            product_id=product_id,
            user_id=session['user_id'],
            rating=rating,
            comment=comment,
            created_at=datetime.now(timezone.utc).replace(tzinfo=None)
        )
        db.session.add(new_review)
        db.session.commit()
        flash("Спасибо! Ваш отзыв успешно опубликован.", "success")
        
    log_system_event("USER", f"User ID {session['user_id']} left a {rating}-star review on Product #{product_id}.")
    return redirect(url_for('product_detail', product_id=product_id))


@app.route('/review/<int:review_id>/reply', methods=['POST'])
def submit_review_reply(review_id):
    """Allows sellers or admins to post replies to user reviews."""
    if 'user_id' not in session:
        flash("Авторизуйтесь, чтобы составить ответ.", "error")
        return redirect(url_for('login'))
        
    review = ProductReview.query.get_or_404(review_id)
    reply_comment = request.form.get('seller_reply', '').strip()
    
    if not reply_comment:
        flash("Вы не можете отправить пустой ответ.", "error")
        return redirect(url_for('product_detail', product_id=review.product_id))
        
    # Check authorization (seller must own the listed product, or be admin)
    if review.product.seller_id != session['user_id'] and session['role'] != 'Admin':
        flash("Вам запрещено отвечать от лица продавца этого товара.", "error")
        return redirect(url_for('product_detail', product_id=review.product_id))
        
    review.seller_reply = reply_comment
    db.session.commit()
    
    flash("Sellers reply successfully published!", "success")
    log_system_event("USER", f"Seller left a reply response on Product Review #{review.id}.")
    return redirect(url_for('product_detail', product_id=review.product_id))


# --- BUYER MODULES (CLIENT DIVISION) ---

@app.route('/buyer/add-to-cart', methods=['POST'])
@role_required(['Buyer', 'Seller', 'Admin'])
def add_to_cart():
    """Pivots products into active temporary 'Cart' orders."""
    product_id = request.form.get('product_id', type=int)
    quantity = request.form.get('quantity', default=1, type=int)
    
    if not quantity or quantity <= 0:
        flash("Invalid quantity count specified.", "error")
        return redirect(request.referrer or url_for('catalog'))
        
    product = Product.query.get_or_404(product_id)
    if product.stock < quantity:
        flash(f"Insufficient inventories. Only {product.stock} units available in stock.", "error")
        return redirect(request.referrer or url_for('catalog'))
        
    # Query current or instantiate new Cart Order
    cart_order = Order.query.filter_by(buyer_id=session['user_id'], status='Cart').first()
    if not cart_order:
        cart_order = Order(buyer_id=session['user_id'], status='Cart', total_price=0.0)
        db.session.add(cart_order)
        db.session.commit()
        
    # Check if Item already mapped in Cart
    cart_item = OrderItem.query.filter_by(order_id=cart_order.id, product_id=product.id).first()
    if cart_item:
        if product.stock < (cart_item.quantity + quantity):
            flash(f"Cannot secure quantity. Max items of this product already in cart.", "error")
            return redirect(url_for('buyer_cart'))
        cart_item.quantity += quantity
    else:
        cart_item = OrderItem(
            order_id=cart_order.id,
            product_id=product.id,
            quantity=quantity,
            price_at_purchase=product.price
        )
        db.session.add(cart_item)
        
    db.session.commit()
    
    # Recalculate whole order pricing
    cart_order.total_price = sum(item.quantity * item.price_at_purchase for item in cart_order.items)
    db.session.commit()
    
    flash(f"Added {quantity} unit(s) of '{product.title}' to your shopping cart.", "success")
    return redirect(url_for('buyer_cart'))


@app.route('/buyer/cart')
@role_required(['Buyer', 'Seller', 'Admin'])
def buyer_cart():
    """Itemized basket with live totaling, removal & modifier hooks."""
    cart_order = Order.query.filter_by(buyer_id=session['user_id'], status='Cart').first()
    return render_template('buyer_cart.html', order=cart_order)


@app.route('/buyer/cart/update/<int:item_id>', methods=['POST'])
@role_required(['Buyer', 'Seller', 'Admin'])
def update_cart_quantity(item_id):
    """Modifies the target item allocation count recursively after inventory assertions."""
    item = OrderItem.query.get_or_404(item_id)
    if item.order.buyer_id != session['user_id'] and session['role'] != 'Admin':
        abort(403)
    action = request.form.get('action') # 'increase' vs 'decrease'
    
    if action == 'increase':
        if item.product.stock > item.quantity:
            item.quantity += 1
        else:
            flash(f"Cannot add more. Vendor only supports {item.product.stock} units.", "error")
    elif action == 'decrease':
        if item.quantity > 1:
            item.quantity -= 1
        else:
            db.session.delete(item)
            
    db.session.commit()
    
    # Reset total
    order = Order.query.get(item.order_id)
    order.total_price = sum(i.quantity * i.price_at_purchase for i in order.items)
    if not order.items:
        db.session.delete(order)
    db.session.commit()
    
    return redirect(url_for('buyer_cart'))


@app.route('/buyer/cart/delete/<int:item_id>', methods=['POST'])
@role_required(['Buyer', 'Seller', 'Admin'])
def delete_cart_item(item_id):
    """Prunes item row from cart."""
    item = OrderItem.query.get_or_404(item_id)
    if item.order.buyer_id != session['user_id'] and session['role'] != 'Admin':
        abort(403)
    order_id = item.order_id
    db.session.delete(item)
    db.session.commit()
    
    # Reset total
    order = Order.query.get(order_id)
    order.total_price = sum(i.quantity * i.price_at_purchase for i in order.items)
    if not order.items:
        db.session.delete(order)
    db.session.commit()
    
    flash("Item successfully removed from shopping cart.", "success")
    return redirect(url_for('buyer_cart'))


@app.route('/buyer/cart/checkout', methods=['POST'])
@role_required(['Buyer', 'Seller', 'Admin'])
def cart_checkout_process():
    """Locks cart order variables, maps destination, prepares checkout page."""
    shipping_address = request.form.get('shipping_address', '').strip()
    if not shipping_address:
        flash("Shipping address is mandatory for invoice routing.", "error")
        return redirect(url_for('buyer_cart'))
        
    cart_order = Order.query.filter_by(buyer_id=session['user_id'], status='Cart').first()
    if not cart_order or not cart_order.items:
        flash("Your shopping cart is empty.", "error")
        return redirect(url_for('catalog'))
        
    # Check physical stock availability for entire set
    for item in cart_order.items:
        if item.product.stock < item.quantity:
            flash(f"Stock conflict detected. Product '{item.product.title}' has only {item.product.stock} units.", "error")
            return redirect(url_for('buyer_cart'))
            
    cart_order.shipping_address = shipping_address
    # Instantly lock purchase prices inside line items
    for item in cart_order.items:
        item.price_at_purchase = item.product.price
        
    cart_order.total_price = sum(item.quantity * item.price_at_purchase for item in cart_order.items)
    db.session.commit()
    
    return redirect(url_for('buyer_checkout', order_id=cart_order.id))


@app.route('/buyer/checkout/<int:order_id>', methods=['GET', 'POST'])
@role_required(['Buyer', 'Seller', 'Admin'])
def buyer_checkout(order_id):
    """Shows direct funding details or file uploader to bridge payment proof."""
    order = Order.query.filter_by(id=order_id, buyer_id=session['user_id']).first_or_404()
    
    # Block bypasses if user already uploaded
    if order.status != 'Cart':
        flash("This order transaction is already submitted or verified.", "info")
        return redirect(url_for('buyer_orders'))
        
    if request.method == 'POST':
        # File parsing
        file = request.files.get('payment_proof')
        if not file or file.filename == '':
            flash("Please upload a valid payment slip files.", "error")
            return redirect(url_for('buyer_checkout', order_id=order.id))
            
        _, ext = os.path.splitext(file.filename)
        if not ext:
            ext = '.png'
        # Unique naming to prevent collisions or Cyrillic drop bugs
        unique_name = f"proof_order_{order.id}_{uuid.uuid4().hex}{ext.lower()}"
        save_path = os.path.join(app.config['UPLOAD_FOLDER'], 'payment_proofs', unique_name)
        file.save(save_path)
        
        # Deduct product stock allocations instantly to prevent double carting, checking stock availability right before lock!
        for item in order.items:
            if item.product.stock < item.quantity:
                flash(f"Inventory conflict detected: '{item.product.title}' has only {item.product.stock} units left on shelf.", "error")
                return redirect(url_for('buyer_checkout', order_id=order.id))
            item.product.stock -= item.quantity
            
        # Overwrite Order records
        order.status = 'Paid/Pending'
        order.payment_proof_path = f"/static/uploads/payment_proofs/{unique_name}"
        db.session.commit()
        
        log_system_event("ORDER", f"Buyer submitted receipt proof {unique_name} for Order ID: {order.id}. Total sum {order.total_price}.")
        flash("Proof of payment transfer successfully uploaded. Pending admin execution validation.", "success")
        return redirect(url_for('buyer_orders'))
        
    return render_template('buyer_checkout.html', order=order)


@app.route('/buyer/checkout/instant/<int:order_id>', methods=['POST'])
@role_required(['Buyer', 'Seller', 'Admin'])
def buyer_checkout_instant(order_id):
    """Processes instant wallet-based payment without manual administrator slip reviews."""
    order = Order.query.filter_by(id=order_id, buyer_id=session['user_id']).first_or_404()
    
    if order.status != 'Cart':
        flash("Order transaction is already finalized or verified.", "info")
        return redirect(url_for('buyer_orders'))
        
    buyer = User.query.get(session['user_id'])
    
    # 1. Validation checks on stock and balance
    for item in order.items:
        if item.product.stock < item.quantity:
            flash(f"Insufficient stock for '{item.product.title}'. Only {item.product.stock} left.", "error")
            return redirect(url_for('buyer_checkout', order_id=order.id))
            
    if buyer.balance < order.total_price:
        flash(f"Sufficient balance not found in wallet. Current balance: ${buyer.balance:.2f}. Total required: ${order.total_price:.2f}.", "error")
        return redirect(url_for('buyer_checkout', order_id=order.id))
        
    try:
        # Deduct balance
        buyer.balance -= order.total_price
        
        # Pay Sellers/Merchants instantly, updating stock cleanly
        for item in order.items:
            # Re-verify stock before physical decrement
            if item.product.stock < item.quantity:
                db.session.rollback()
                flash(f"Stock conflict detected mid-transaction for '{item.product.title}'.", "error")
                return redirect(url_for('buyer_checkout', order_id=order.id))
                
            item.product.stock -= item.quantity
            merchant = User.query.get(item.product.seller_id)
            if merchant:
                transaction_split = item.quantity * item.price_at_purchase
                merchant.balance += transaction_split
                log_system_event("ORDER", f"Credited merchant {merchant.full_name} with ${transaction_split} for product {item.product.title} (Instant Payment).")
                
        order.status = 'Completed' # Instant checkout automatically fulfills order
        db.session.commit()
        
        log_system_event("ORDER", f"Buyer {buyer.email} completed INSTANT WALLET CHECKOUT for Order ID {order.id}. Sum: ${order.total_price}.")
        flash(f"Payment successful! ${order.total_price:.2f} processed. Your items are now available and your order is Completed.", "success")
        return redirect(url_for('buyer_orders'))
        
    except Exception as e:
        db.session.rollback()
        log_system_event("SYSTEM", f"Critical exception during instant checkout for Order {order_id}: {str(e)}")
        flash("A server error interrupted your instant checkout transaction. Please try again.", "error")
        return redirect(url_for('buyer_checkout', order_id=order.id))


@app.route('/buyer/orders')
@role_required(['Buyer', 'Seller', 'Admin'])
def buyer_orders():
    """Client purchase log displaying real-time colored badge states."""
    orders = Order.query.filter_by(buyer_id=session['user_id']).order_by(Order.id.desc()).all()
    # Filter active cart orders out
    orders_filtered = [o for o in orders if o.status != 'Cart']
    return render_template('buyer_orders.html', orders=orders_filtered)


# --- SELLER MODULES (MERCHANT DIVISION) ---

@app.route('/seller/dashboard')
@role_required(['Seller', 'Admin'])
def seller_dashboard():
    """Dynamic Sales dashboards illustrating KPIs and active pipelines."""
    seller_id = session['user_id']
    
    # 1. Active metrics
    products_count = Product.query.filter_by(seller_id=seller_id, is_active=True).count()
    
    # 2. Get all orders containing products belonging to this seller
    all_seller_order_items = OrderItem.query.join(Product).filter(Product.seller_id == seller_id).all()
    
    total_earnings = 0.0
    pending_fulfillments_count = 0
    completed_shipments_count = 0
    
    purchased_items = []
    
    for item in all_seller_order_items:
        o = item.order
        if o.status in ['Paid/Pending', 'Shipped', 'Completed']:
            item_sum = item.quantity * item.price_at_purchase
            if o.status == 'Completed':
                total_earnings += item_sum
                completed_shipments_count += 1
            elif o.status == 'Paid/Pending':
                pending_fulfillments_count += 1
                
            purchased_items.append({
                'order_id': o.id,
                'buyer_name': o.buyer.full_name,
                'product_title': item.product.title,
                'quantity': item.quantity,
                'item_total': item_sum,
                'shipping_address': o.shipping_address,
                'status': o.status,
                'created_at': o.created_at
            })
            
    purchased_items.sort(key=lambda x: x['created_at'], reverse=True)
            
    return render_template('seller_dashboard.html', 
                           products_count=products_count, 
                           total_earnings=total_earnings, 
                           pending_fulfillments=pending_fulfillments_count,
                           completed_shipments=completed_shipments_count,
                           recent_sales=purchased_items[:8])


@app.route('/seller/products', methods=['GET', 'POST'])
@role_required(['Seller', 'Admin'])
def seller_products():
    """Full inventory listing and modal add forms."""
    seller_id = session['user_id']
    products = Product.query.filter_by(seller_id=seller_id).order_by(Product.id.desc()).all()
    return render_template('seller_products.html', products=products)


@app.route('/seller/products/add', methods=['POST'])
@role_required(['Seller', 'Admin'])
def seller_add_product():
    """Inserts a new listing mapping extended details dynamically to metadata_json."""
    title = request.form.get('title', '').strip()
    description = request.form.get('description', '').strip()
    category = request.form.get('category', '').strip()
    price = request.form.get('price', type=float)
    stock = request.form.get('stock', type=int)
    image_url = request.form.get('image_url', '').strip()
    
    # Custom attributes from UI form
    attr_keys = request.form.getlist('attr_keys[]')
    attr_vals = request.form.getlist('attr_vals[]')
    
    # Zip custom fields into python dict
    custom_metadata = {}
    for k, v in zip(attr_keys, attr_vals):
        if k.strip():
            custom_metadata[k.strip()] = v.strip()
            
    if not title or not description or not category or price is None or stock is None:
        flash("Please fill in all core product information requirements.", "error")
        return redirect(url_for('seller_products'))
        
    if price < 0.0 or stock < 0:
        flash("Price and stock values cannot be negative.", "error")
        return redirect(url_for('seller_products'))
        
    if not image_url:
        image_url = "https://images.unsplash.com/photo-1531403009284-440f080d1e12?auto=format&fit=crop&w=600&q=80"
        
    new_prod = Product(
        seller_id=session['user_id'],
        title=title,
        description=description,
        category=category,
        price=price,
        stock=stock,
        image_url=image_url,
        is_active=True
    )
    new_prod.set_metadata(custom_metadata)
    
    db.session.add(new_prod)
    db.session.commit()
    
    log_system_event("CSV_IMPORT" if "imported" in title.lower() else "SYSTEM", f"Seller added new catalog product title '{title}' in category {category}.")
    flash(f"Catalog product '{title}' successfully generated.", "success")
    return redirect(url_for('seller_products'))


@app.route('/seller/products/edit/<int:product_id>', methods=['POST'])
@role_required(['Seller', 'Admin'])
def seller_edit_product(product_id):
    """Updates existing product details including custom metadata attributes list."""
    product = Product.query.get_or_404(product_id)
    # Protection rules assert owner
    if product.seller_id != session['user_id'] and session['role'] != 'Admin':
        abort(403)
        
    title_val = request.form.get('title', '').strip()
    desc_val = request.form.get('description', '').strip()
    cat_val = request.form.get('category', '').strip()
    price_val = request.form.get('price', type=float)
    stock_val = request.form.get('stock', type=int)
    image_val = request.form.get('image_url', '').strip()
    
    if not title_val or not desc_val or not cat_val or price_val is None or stock_val is None:
        flash("Please fill in all core product details during modification.", "error")
        return redirect(url_for('seller_products'))
        
    if price_val < 0.0 or stock_val < 0:
        flash("Price and stock values cannot be negative.", "error")
        return redirect(url_for('seller_products'))
        
    product.title = title_val
    product.description = desc_val
    product.category = cat_val
    product.price = price_val
    product.stock = stock_val
    product.image_url = image_val
    
    # Re-evaluate custom variables
    attr_keys = request.form.getlist('attr_keys[]')
    attr_vals = request.form.getlist('attr_vals[]')
    custom_metadata = {}
    for k, v in zip(attr_keys, attr_vals):
        if k.strip():
            custom_metadata[k.strip()] = v.strip()
            
    product.set_metadata(custom_metadata)
    db.session.commit()
    
    log_system_event("SYSTEM", f"Sellers modified details mapping product #{product.id} : {product.title}")
    flash(f"Product '{product.title}' successfully modified.", "success")
    return redirect(url_for('seller_products'))


@app.route('/seller/products/delete/<int:product_id>', methods=['POST'])
@role_required(['Seller', 'Admin'])
def seller_delete_product(product_id):
    """Delete a listing from display list."""
    product = Product.query.get_or_404(product_id)
    if product.seller_id != session['user_id'] and session['role'] != 'Admin':
        abort(403)
        
    title = product.title
    db.session.delete(product)
    db.session.commit()
    
    log_system_event("SYSTEM", f"Inventory product ID {product_id} deleted by seller.")
    flash(f"Product '{title}' removed from active inventory.", "success")
    return redirect(url_for('seller_products'))


@app.route('/seller/orders')
@role_required(['Seller', 'Admin'])
def seller_orders():
    """Merchant order fulfillment tracker showing delivery address details."""
    seller_id = session['user_id']
    all_seller_items = OrderItem.query.join(Product).filter(Product.seller_id == seller_id).all()
    
    # Pivot rows around unique orders
    orders_map = {}
    for item in all_seller_items:
        o = item.order
        if o.status in ['Paid/Pending', 'Shipped', 'Completed', 'Cancelled']:
            if o.id not in orders_map:
                orders_map[o.id] = {
                    'order_id': o.id,
                    'buyer_name': o.buyer.full_name,
                    'shipping_address': o.shipping_address,
                    'status': o.status,
                    'created_at': o.created_at,
                    'products': []
                }
            orders_map[o.id]['products'].append({
                'title': item.product.title,
                'quantity': item.quantity,
                'price': item.price_at_purchase,
                'total': item.quantity * item.price_at_purchase
            })
            
    orders_list = list(orders_map.values())
    orders_list.sort(key=lambda x: x['created_at'], reverse=True)
    return render_template('seller_orders.html', orders=orders_list)


@app.route('/seller/orders/ship/<int:order_id>', methods=['POST'])
@role_required(['Seller', 'Admin'])
def seller_ship_order(order_id):
    """Fulfillment state transition handler update."""
    order = Order.query.get_or_404(order_id)
    
    # Assert containment
    seller_id = session['user_id']
    owns_product = False
    for item in order.items:
        if item.product.seller_id == seller_id or session['role'] == 'Admin':
            owns_product = True
            break
            
    if not owns_product:
        abort(403)
        
    if order.status == 'Paid/Pending':
        order.status = 'Shipped'
        db.session.commit()
        log_system_event("ORDER", f"Merchant flagged Order ID {order.id} as 'Shipped'. Delivery route initialized.")
        flash("Order marked as Shipped. Customer notified via order ledger.", "success")
    else:
        flash("Order cannot transition to Shipped from its current state.", "error")
        
    return redirect(url_for('seller_orders'))


@app.route('/seller/import', methods=['GET', 'POST'])
@role_required(['Seller', 'Admin'])
def seller_bulk_import():
    """Bulk Inventory File processing engine parses CSV or JSON arrays uploads."""
    if request.method == 'POST':
        file = request.files.get('import_file')
        if not file or file.filename == '':
            flash("Please choose a valid JSON or CSV spreadsheet format.", "error")
            return redirect(url_for('seller_bulk_import'))
            
        _, ext = os.path.splitext(file.filename)
        if not ext:
            ext = '.csv' if 'csv' in file.filename.lower() else '.json'
        filename = f"{uuid.uuid4().hex}{ext.lower()}"
        save_path = os.path.join(app.config['UPLOAD_FOLDER'], 'bulk_imports', filename)
        file.save(save_path)
        
        imported_count = 0
        
        try:
            # Parse CSV Logic
            if filename.endswith('.csv'):
                with open(save_path, mode='r', encoding='utf-8') as f:
                    reader = csv.DictReader(f)
                    for row in reader:
                        # Extract row parameters
                        title = row.get('title', '').strip()
                        description = row.get('description', '').strip()
                        category = row.get('category', '').strip()
                        price = float(row.get('price', 0.0))
                        stock = int(row.get('stock', 1))
                        image_url = row.get('image_url', '').strip()
                        
                        # Parsing dynamic variables
                        metadata_raw = row.get('metadata', '{}')
                        try:
                            meta_dict = json.loads(metadata_raw)
                        except Exception:
                            meta_dict = {"source": "CSV Bulk Import Detail"}
                            
                        if title and category:
                            new_p = Product(
                                seller_id=session['user_id'],
                                title=title,
                                description=description or "Standard imported catalog item template.",
                                category=category,
                                price=price,
                                stock=stock,
                                image_url=image_url or "https://images.unsplash.com/photo-1531403009284-440f080d1e12?auto=format&fit=crop&w=600&q=80",
                                is_active=True
                            )
                            new_p.set_metadata(meta_dict)
                            db.session.add(new_p)
                            imported_count += 1
                db.session.commit()
                
            # Parse JSON Logic    
            elif filename.endswith('.json'):
                with open(save_path, 'r', encoding='utf-8') as f:
                    p_list = json.load(f)
                    if isinstance(p_list, dict):
                        # Wrap into list if single product structure
                        p_list = [p_list]
                        
                    for row in p_list:
                        title = row.get('title', '').strip()
                        category = row.get('category', '').strip()
                        if title and category:
                            new_p = Product(
                                seller_id=session['user_id'],
                                title=title,
                                description=row.get('description', 'No details supplied. Please populate details.'),
                                category=category,
                                price=float(row.get('price', 0.0)),
                                stock=int(row.get('stock', 1)),
                                image_url=row.get('image_url', "https://images.unsplash.com/photo-1531403009284-440f080d1e12?auto=format&fit=crop&w=600&q=80"),
                                is_active=True
                            )
                            new_p.set_metadata(row.get('metadata', {}))
                            db.session.add(new_p)
                            imported_count += 1
                db.session.commit()
                
            else:
                flash("Unsupported parser configuration. Only spreadsheet standards CSV or JSON accepted.", "error")
                return redirect(url_for('seller_bulk_import'))
                
            log_system_event("CSV_IMPORT", f"Processed bulk spreadsheets. Imported {imported_count} listings successfully.")
            flash(f"Bulk operations complete! Added {imported_count} new product listings to your inventory.", "success")
            return redirect(url_for('seller_products'))
            
        except Exception as e:
            db.session.rollback()
            flash(f"Operational breakdown parsing structure: {str(e)}", "error")
            return redirect(url_for('seller_bulk_import'))
            
    return render_template('seller_import.html')


# --- ADMIN DIVISION (MARKETPLACE HEADQUARTERS) ---

@app.route('/admin/dashboard')
@role_required(['Admin'])
def admin_dashboard():
    """Analytic nerve center aggregating platform total GMVs, balances, uploads, and logs."""
    # 1. Total platform GMV calculation (Completed/Shipped/Paid states)
    gmv_orders = Order.query.filter(Order.status.in_(['Paid/Pending', 'Shipped', 'Completed'])).all()
    gmv_sum = sum(o.total_price for o in gmv_orders)
    
    # 2. Total active users on platform
    total_users = User.query.count()
    
    # 3. Active pending transactions uploading receipts
    pending_receipts = Order.query.filter_by(status='Paid/Pending').count()
    
    # 4. Total audit actions
    total_logs = SystemLog.query.count()
    
    # Fetch log items list
    recent_logs = SystemLog.query.order_by(SystemLog.id.desc()).limit(15).all()
    
    return render_template('admin_dashboard.html', 
                           gmv=gmv_sum, 
                           users=total_users, 
                           pending_verifications=pending_receipts, 
                           logs_count=total_logs,
                           recent_logs=recent_logs)


@app.route('/admin/roles', methods=['GET', 'POST'])
@role_required(['Admin'])
def admin_roles():
    """Admin privilege ledger listing and permission overrides."""
    users = User.query.order_by(User.id.asc()).all()
    
    if request.method == 'POST':
        target_user_id = request.form.get('user_id', type=int)
        new_role = request.form.get('role').strip()
        new_balance = request.form.get('balance', type=float)
        
        target_user = User.query.get_or_404(target_user_id)
        
        # Enforce administrative preservation safeguarding
        if target_user.id == session['user_id'] and new_role != 'Admin':
            flash("Self-demotion error. Administrative protection locked.", "error")
            return redirect(url_for('admin_roles'))
            
        target_user.role = new_role
        if new_balance is not None:
            target_user.balance = new_balance
            
        db.session.commit()
        log_system_event("USER", f"Admin reconfigured User ID {target_user.id} ({target_user.email}) Permissions to: Role={new_role}, Balance=${new_balance}.")
        flash(f"System permissions updated for user '{target_user.full_name}'.", "success")
        return redirect(url_for('admin_roles'))
        
    return render_template('admin_roles.html', all_users=users)


@app.route('/admin/verifications')
@role_required(['Admin'])
def admin_verifications():
    """List and manage buyer's uploaded receipts to finalize database billing updates."""
    pending_orders = Order.query.filter_by(status='Paid/Pending').order_by(Order.id.asc()).all()
    return render_template('admin_verifications.html', orders=pending_orders)


@app.route('/admin/verifications/approve/<int:order_id>', methods=['POST'])
@role_required(['Admin'])
def approve_order_payment(order_id):
    """Verifies cash slips, transfers buyer balance, and transitions orders to Paid or Completed."""
    order = Order.query.get_or_404(order_id)
    if order.status != 'Paid/Pending':
        flash("Order doesn't require payment confirmation action.", "error")
        return redirect(url_for('admin_verifications'))
        
    # Standard transaction verification mapping
    # Buyer balance is checked/deducted during manual slip approval, or wallet allocation processed
    buyer = User.query.get(order.buyer_id)
    
    # Since manual invoice transfer represents real Cash wired into our bank account:
    # We credit the buyer's wallet with the verified cash transfer amount first, 
    # then debit it to clear the order's invoice. This keeps their balance non-negative and robust.
    incoming_cash = order.total_price
    buyer.balance += incoming_cash
    log_system_event("ORDER", f"Credited buyer {buyer.full_name} with ${incoming_cash:.2f} due to verified Bank Wire (Order {order.id}).")
    
    # Debit target buyer invoice
    buyer.balance -= order.total_price
    
    # Pay Sellers of corresponding catalog items
    for item in order.items:
        merchant = User.query.get(item.product.seller_id)
        if merchant:
            transaction_split = item.quantity * item.price_at_purchase
            merchant.balance += transaction_split
            log_system_event("ORDER", f"Credited merchant {merchant.full_name} with ${transaction_split} for product {item.product.title}.")
            
    order.status = 'Completed' # Once fundings clear, move order to final dispatch
    db.session.commit()
    
    log_system_event("ORDER", f"Admin APPROVED Payment Verification for Order ID {order.id}. Account ledger finalized.")
    flash(f"Payment slip confirmed! Order ID {order.id} verified and marked as Completed.", "success")
    return redirect(url_for('admin_verifications'))


@app.route('/admin/verifications/decline/<int:order_id>', methods=['POST'])
@role_required(['Admin'])
def decline_order_payment(order_id):
    """Reverts product inventories and moves status to Cancelled."""
    order = Order.query.get_or_404(order_id)
    if order.status != 'Paid/Pending':
        flash("Order context invalid for cancellations.", "error")
        return redirect(url_for('admin_verifications'))
        
    # Restore stock quantities back to catalog items
    for item in order.items:
        item.product.stock += item.quantity
        
    order.status = 'Cancelled'
    db.session.commit()
    
    log_system_event("ORDER", f"Admin REJECTED Payment Slip for Order ID {order.id}. Inventory allocations restored.")
    flash(f"Payment slip declined. Order ID {order.id} transitioned to Cancelled.", "info")
    return redirect(url_for('admin_verifications'))


@app.route('/admin/export-reports')
@role_required(['Admin'])
def admin_export_reports():
    """Generates complete spreadsheet/document exports across all standard judge formats."""
    export_format = request.args.get('format', 'json').lower()
    
    users = User.query.all()
    products = Product.query.all()
    orders = Order.query.all()
    
    title_report = "Official Hackathon Judging Transaction Log"
    timestamp_str = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    
    # 1. JSON Export
    if export_format == 'json':
        report_data = {
            "report_title": title_report,
            "generated_at": timestamp_str,
            "summary": {
                "total_users": len(users),
                "total_catalog_products": len(products),
                "total_orders_placed": len(orders),
                "system_gmv": sum(o.total_price for o in orders if o.status in ['Completed', 'Shipped', 'Paid/Pending'])
            },
            "users": [u.to_dict() for u in users],
            "products": [p.to_dict() for p in products],
            "orders": [o.to_dict() for o in orders]
        }
        
        buffer = io.BytesIO()
        buffer.write(json.dumps(report_data, indent=2).encode('utf-8'))
        buffer.seek(0)
        return send_file(
            buffer,
            mimetype='application/json',
            as_attachment=True,
            download_name=f"marketplace_report_{int(datetime.now(timezone.utc).timestamp())}.json"
        )
        
    # 2. CSV Export
    elif export_format == 'csv':
        buffer = io.StringIO()
        writer = csv.writer(buffer)
        
        # Meta row
        writer.writerow([title_report])
        writer.writerow(["Generated at: " + timestamp_str])
        writer.writerow([])
        
        # User details table
        writer.writerow(["--- USER ROSTER ---"])
        writer.writerow(["User ID", "Full Name", "Email", "Role", "Wallet Balance", "Created At"])
        for u in users:
            writer.writerow([u.id, u.full_name, u.email, u.role, u.balance, u.created_at])
            
        writer.writerow([])
        # Product details table
        writer.writerow(["--- ACTIVE PRODUCTS CATALOG ---"])
        writer.writerow(["Product ID", "Seller ID", "Title", "Category", "Price", "Stock Level", "Active Status"])
        for p in products:
            writer.writerow([p.id, p.seller_id, p.title, p.category, p.price, p.stock, p.is_active])
            
        writer.writerow([])
        # Orders table
        writer.writerow(["--- TRANSACTION ORDERS LEDGER ---"])
        writer.writerow(["Order ID", "Buyer ID", "Total Price", "Status", "Address"])
        for o in orders:
            writer.writerow([o.id, o.buyer_id, o.total_price, o.status, o.shipping_address])
            
        byte_buffer = io.BytesIO()
        byte_buffer.write(buffer.getvalue().encode('utf-8'))
        byte_buffer.seek(0)
        return send_file(
            byte_buffer,
            mimetype='text/csv',
            as_attachment=True,
            download_name=f"marketplace_report_{int(datetime.now(timezone.utc).timestamp())}.csv"
        )
        
    # 3. TEXT Export
    elif export_format == 'txt':
        buffer = io.StringIO()
        buffer.write(f"==========================================================\n")
        buffer.write(f" {title_report.upper()}\n")
        buffer.write(f"==========================================================\n")
        buffer.write(f"Generated at: {timestamp_str}\n\n")
        
        buffer.write(f"--- PLATFORM KPIs SUMMARY ---\n")
        buffer.write(f" Total Registered Users: {len(users)}\n")
        buffer.write(f" Active Product Listings: {len(products)}\n")
        buffer.write(f" Total Orders Handled: {len(orders)}\n")
        active_gmv = sum(o.total_price for o in orders if o.status in ['Completed', 'Shipped', 'Paid/Pending'])
        buffer.write(f" Gross Merchandise Value (GMV): ${active_gmv:.2f}\n\n")
        
        buffer.write(f"--- PLATFORM USERS INVENTORY ---\n")
        for u in users:
            buffer.write(f" [{u.role.upper()}] ID#{u.id}: {u.full_name} | {u.email} | Bal: ${u.balance:.2f}\n")
            
        buffer.write(f"\n--- CATALOG LISTINGS DETAILS ---\n")
        for p in products:
            buffer.write(f" ID#{p.id}: {p.title} | Cat: {p.category} | Price: ${p.price:.2f} | Stock: {p.stock} units\n")
            
        buffer.write(f"\n--- ORDER TRANSACTIONS PROCESS ---\n")
        for o in orders:
            buyer_user = User.query.get(o.buyer_id)
            buyer_name = buyer_user.full_name if buyer_user else "Deleted Buyer"
            buffer.write(f" ORDER #{o.id} - Buyer: {buyer_name} | Total: ${o.total_price:.2f} | Status: {o.status.upper()} | Dest: {o.shipping_address}\n")
            
        byte_buffer = io.BytesIO()
        byte_buffer.write(buffer.getvalue().encode('utf-8'))
        byte_buffer.seek(0)
        return send_file(
            byte_buffer,
            mimetype='text/plain',
            as_attachment=True,
            download_name=f"marketplace_report_{int(datetime.now(timezone.utc).timestamp())}.txt"
        )
        
    # 4. DOCX Word Document Export (using python-docx)
    elif export_format == 'docx':
        doc = Document()
        
        doc.add_heading(title_report, 0)
        doc.add_paragraph(f"Generated on official sandbox servers: {timestamp_str}")
        
        doc.add_heading('Platform Performance Matrix', level=1)
        doc.add_paragraph(f"Active Users: {len(users)}\nTotal Product Listings: {len(products)}\nTotal Transactions Audited: {len(orders)}")
        
        # Add summary table
        p_table = doc.add_table(rows=1, cols=4)
        hdr_cells = p_table.rows[0].cells
        hdr_cells[0].text = 'Metric Code'
        hdr_cells[1].text = 'Platform Scope'
        hdr_cells[2].text = 'Financial State'
        hdr_cells[3].text = 'Active Ledger Status'
        
        # Seed simple matrix table values
        row_cells = p_table.add_row().cells
        row_cells[0].text = 'GMV_VAL'
        row_cells[1].text = 'Global marketplace transactions'
        row_cells[2].text = f"${sum(o.total_price for o in orders if o.status in ['Completed', 'Shipped', 'Paid/Pending']):.2f}"
        row_cells[3].text = 'Approved and pending clearance'
        
        # Detailed audit blocks
        doc.add_heading('Core User Account Register', level=1)
        for u in users:
            doc.add_paragraph(f"User ID: {u.id} | Name: {u.full_name} | Role: {u.role} | Wallet: ${u.balance:.2f} | Email: {u.email}", style='List Bullet')
            
        doc.add_heading('Catalog Inventory Allocation', level=1)
        for p in products:
            doc.add_paragraph(f"Listing ID: {p.id} | Title: {p.title} | Category: {p.category} | Price: ${p.price:.2f} | Stock: {p.stock}", style='List Bullet')
            
        doc.add_heading('Transactions Accounting History', level=1)
        for o in orders:
            doc.add_paragraph(f"Order ID: {o.id} | Total Sum: ${o.total_price:.2f} | Status: {o.status} | Shipping Addr: {o.shipping_address}", style='List Bullet')
            
        file_stream = io.BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)
        return send_file(
            file_stream,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            as_attachment=True,
            download_name=f"marketplace_report_{int(datetime.now(timezone.utc).timestamp())}.docx"
        )
        
    else:
        abort(400)


if __name__ == '__main__':
    # Bind to host 0.0.0.0 and port 3000 as structured in standard environment requirements
    app.run(host='0.0.0.0', port=3000, debug=True)
